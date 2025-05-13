from fastapi import FastAPI, File, UploadFile
from typing import List
from transformers import pipeline
import os
import uuid
from docx import Document
import fitz  # PyMuPDF for PDFs
import pytesseract
from PIL import Image
import io
from sqlalchemy.orm import Session
from database import Session, Document

app = FastAPI()
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Load the Hugging Face summarization pipeline
summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")
print("Model downloaded successfully!")

# --- Text Extraction Functions ---

def extract_text_from_pdf(file_bytes):
    text = ""
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(file_bytes):
    document = Document(io.BytesIO(file_bytes))
    return "\n".join([para.text for para in document.paragraphs])

def extract_text_from_image(file_bytes):
    image = Image.open(io.BytesIO(file_bytes))
    return pytesseract.image_to_string(image)


@app.get("/")
def read_root():
    return {"message": "SmartDoc AI is running"}
@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    file_id = str(uuid.uuid4())
    file_path = os.path.join(UPLOAD_DIR, f"{file_id}_{file.filename}")
    
    
    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())
    
    # Save file metadata to the database
    session = Session()
    try:
        new_document = Document(
            id=file_id,
            filename=file.filename,
            file_path=file_path,
            file_type=file.content_type
        )
        session.add(new_document)
        session.commit()
    finally:
        session.close()
    
    return {"file_id": file_id, "filename": file.filename, "summary": "File uploaded successfully"}
@app.get("/files")
def get_files():
    session = Session()
    try:
        documents = session.query(Document).all()
        files = [{"id": doc.id, "filename": doc.filename} for doc in documents]
    finally:
        session.close()
    
    return {"files": files} 

# --- Summarization Endpoint ---
@app.post("/summarize")
async def summarize_file(file: UploadFile = File(...)):
    file_bytes = await file.read()
    filename = file.filename.lower()

    if filename.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    elif filename.endswith((".png", ".jpg", ".jpeg")):
        text = extract_text_from_image(file_bytes)
    else:
        return {"error": "Unsupported file format."}

    if not text.strip():
        return {"error": "No text found in the document."}

    # Limit the text size for summarization (Hugging Face limit)
    max_chunk = 2048  # characters
    text = text[:max_chunk]

    summary = summarizer(text, max_length=400, min_length=150, do_sample=False)

    return {"summary": summary[0]['summary_text']}
